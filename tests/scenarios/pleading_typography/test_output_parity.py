"""Output parity for md_to_docx.py and md_to_txt.py (spec:
specs/pleading/generator.md; feature inventory:
pleading/pleading_markdown_spec.md).

The DOCX and TXT renderers share md_pleading's front-end (inline
parsing, auto-numbering, footnote extraction) and must carry the same
content as the PDF for the same source. These checks open the actual
artifacts — the .docx via python-docx (proving it is a valid OPC
package Word can load), the .txt as structured text — and never trust
the converters' self-report.

xfail'd tests are real parity bugs — see
design/refactor-audit/typography_structure.md.
"""

from __future__ import annotations

import datetime
import subprocess
import sys

import pytest

from tests.harness import scenario

DECL_MD = "Declaration of Jane Roe.md"
MD_TO_TXT = scenario.PLEADING / "md_to_txt.py"
MD_TO_DOCX = scenario.PLEADING / "md_to_docx.py"


@pytest.fixture(scope="module")
def built(tmp_path_factory):
    matter = scenario.load_scenario("pleading_typography",
                                    tmp_path_factory.mktemp("m"))
    proc = scenario.build_envelope(matter, "typography")
    assert proc.returncode == 0, proc.stderr[-2000:]
    return matter


@pytest.fixture(scope="module")
def docx_doc(built):
    """The envelope's docx: true output, opened for real by python-docx."""
    from docx import Document
    path = built / "out" / "typography" / "Declaration of Jane Roe.docx"
    assert path.exists(), "docx: true source produced no .docx"
    return Document(str(path))


@pytest.fixture(scope="module")
def docx_paras(docx_doc):
    return [p.text for p in docx_doc.paragraphs]


@pytest.fixture(scope="module")
def txt_out(built, tmp_path_factory):
    out = tmp_path_factory.mktemp("txt") / "decl.txt"
    proc = subprocess.run(
        [sys.executable, str(MD_TO_TXT),
         str(built / "src" / DECL_MD), str(out)],
        capture_output=True, text=True)
    assert proc.returncode == 0, proc.stderr[-1000:]
    return out.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# DOCX: typography
# ---------------------------------------------------------------------------

def test_docx_dashes_converted_and_unspaced(docx_paras):
    joined = "\n".join(docx_paras)
    assert "January—TKEMD1a—and" in joined
    assert "January 23–March 3, 2026" in joined
    assert "---" not in joined and "--" not in joined
    assert " — " not in joined and " – " not in joined


def test_docx_keeps_straight_quotes_by_design(docx_paras):
    """DOCX house style (md_to_docx.typographic_subs): dashes convert,
    quotes/apostrophes stay ASCII so Word's own autocorrect can manage
    them. Pinned so a silent change to either half is caught."""
    joined = "\n".join(docx_paras)
    assert '"unworkable"' in joined
    assert "Ms. Roe's counsel" in joined
    assert "“" not in joined and "’" not in joined


def test_docx_section_symbols(docx_paras):
    joined = "\n".join(docx_paras)
    assert "§ 2030.300" in joined and "¶ 4" in joined


# ---------------------------------------------------------------------------
# DOCX: structure (headings, footnotes, signature block)
# ---------------------------------------------------------------------------

# heading text (docx carries no literal numeral) -> expected w:ilvl
DOCX_HEADINGS = [
    ("INTRODUCTION", 0),
    ("FACTUAL BACKGROUND", 0),
    ("The Records Requests", 1),
    ("The First Email", 2),
    ("The Second Email", 2),
    ("The Meet-and-Confer Process", 1),
    ("Early Calls", 2),
    ("ARGUMENT", 0),
    ("Legal Standard", 1),
    ("Application", 1),
]


def _numbered_paragraphs(doc):
    from docx.oxml.ns import qn
    out = []
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue
        ilvl = int(numPr.find(qn("w:ilvl")).get(qn("w:val")))
        out.append((p.text, ilvl))
    return out


def test_docx_headings_use_word_native_multilevel_numbering(docx_doc):
    assert _numbered_paragraphs(docx_doc) == DOCX_HEADINGS


def test_docx_headings_carry_no_literal_numerals(docx_paras):
    assert not any(t.startswith(("I. ", "II. ", "III. ")) for t in docx_paras), (
        "a literal outline numeral leaked into a docx heading — it would "
        "double up with Word's own numbering")


def test_docx_numbering_definition_is_legal_outline(docx_doc):
    xml = docx_doc.part.numbering_part.element.xml
    for fmt in ("upperRoman", "upperLetter", "decimal"):
        assert fmt in xml, f"multilevel numbering lacks {fmt} level"


def test_docx_footnotes_are_real_word_footnotes(docx_doc):
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    rels = [r for r in docx_doc.part.rels.values() if r.reltype == RT.FOOTNOTES]
    assert rels, "no footnotes part — notes were dropped or inlined"
    notes_xml = rels[0].target_part.blob.decode("utf-8")
    for tok in ("TKFN1", "TKFN2", "TKFN3"):
        assert tok in notes_xml, f"footnote text {tok} missing from part"
    assert "<w:i/>" in notes_xml, "italic run inside a footnote lost"
    assert "footnoteReference" in docx_doc.element.xml, (
        "body carries no footnote references")


def test_docx_declsignblock(docx_doc, docx_paras):
    year = datetime.date.today().year
    executed = (f"Executed this _____ day of _________________, {year}, "
                "at Springfield, California.")
    assert executed in docx_paras
    assert "____________________________________" in docx_paras
    assert "JANE ROE" in docx_paras
    assert "Declarant, In Pro Se" in docx_paras
    # The block must be marked keep-with-next so Word cannot split it.
    for p in docx_doc.paragraphs:
        if p.text == executed:
            assert p.paragraph_format.keep_with_next, (
                "signature block not kept together in Word")
            break


def test_docx_caption_dashes_converted(built, tmp_path):
    """Caption metadata (YAML display strings) gets the same dash
    substitutions as body text in the DOCX, matching the PDF caption
    (spec: pleading_markdown_spec.md, typographic substitutions)."""
    src = tmp_path / "capdash.md"
    src.write_text(
        (built / "src" / DECL_MD).read_text(encoding="utf-8").replace(
            'paper_title: "DECLARATION OF JANE ROE IN SUPPORT OF MOTION '
            'TO COMPEL"',
            'paper_title: "DECLARATION RE JANUARY 23--MARCH 3, 2026 '
            'LOG---TKCAP1"'),
        encoding="utf-8")
    out = tmp_path / "capdash.docx"
    proc = subprocess.run(
        [sys.executable, str(MD_TO_DOCX), str(src), str(out)],
        capture_output=True, text=True)
    assert proc.returncode == 0, proc.stderr[-1000:]
    from docx import Document
    doc = Document(str(out))
    cells = "\n".join(c.text for t in doc.tables for r in t.rows
                      for c in r.cells)
    assert "23–MARCH 3" in cells, "en dash not converted in docx caption"
    assert "LOG—TKCAP1" in cells, "em dash not converted in docx caption"
    assert "--" not in cells and "---" not in cells


def test_docx_attachment_macro_parity(built, tmp_path):
    src = tmp_path / "attach.md"
    src.write_text(
        (built / "src" / DECL_MD).read_text(encoding="utf-8").replace(
            "I, Jane Roe, declare as follows:",
            "I, Jane Roe, declare as follows: a true and correct copy of "
            "the log is attached as \\attachment{log} (TKATT1).\n\n"
            "Original opener follows."),
        encoding="utf-8")
    # Give the source an exhibits list so the shortname resolves.
    text = src.read_text(encoding="utf-8")
    text = text.replace(
        "paper_title: \"DECLARATION OF JANE ROE IN SUPPORT OF MOTION TO "
        "COMPEL\"",
        "paper_title: \"DECLARATION OF JANE ROE IN SUPPORT OF MOTION TO "
        "COMPEL\"\nexhibits:\n  - shortname: \"log\"\n    title: "
        "\"Exchange Log\"\n    path: \"log.pdf\"")
    src.write_text(text, encoding="utf-8")
    out = tmp_path / "attach.docx"
    proc = subprocess.run(
        [sys.executable, str(MD_TO_DOCX), str(src), str(out)],
        capture_output=True, text=True)
    assert proc.returncode == 0, proc.stderr[-1000:]
    from docx import Document
    joined = "\n".join(p.text for p in Document(str(out)).paragraphs)
    line = next(l for l in joined.splitlines() if "TKATT1" in l)
    assert "\\attachment" not in line and "Exhibit A" in line


# ---------------------------------------------------------------------------
# TXT: typography and structure
# ---------------------------------------------------------------------------

def test_txt_typography_matches_pdf(txt_out):
    assert "January—TKEMD1a—and" in txt_out
    assert "January 23–March 3, 2026" in txt_out
    assert "“That schedule—the one we signed—is ‘final,’ period.”" in txt_out
    assert "Ms. Roe’s counsel" in txt_out
    assert "§ 2030.300" in txt_out and "¶ 4" in txt_out
    assert " — " not in txt_out and " --- " not in txt_out


def test_txt_repairs_possessive_after_abbreviation(txt_out):
    """The after-abbreviation-period apostrophe is handled in the shared
    typographic_subs (md_pleading.py), so the TXT and PDF artifacts agree —
    the PDF-side counterpart is test_possessive_after_abbreviation_period
    in test_typography.py."""
    assert "C.E.O.’s calendar" in txt_out
    assert "C.E.O.‘s" not in txt_out


def test_txt_heading_numbering_literal(txt_out):
    pos = 0
    for needle in ["I. INTRODUCTION", "II. FACTUAL BACKGROUND",
                   "A. The Records Requests", "1. The First Email",
                   "B. The Meet-and-Confer Process", "III. ARGUMENT",
                   "A. Legal Standard"]:
        idx = txt_out.find(needle, pos)
        assert idx >= 0, f"txt heading {needle!r} missing or out of order"
        pos = idx + len(needle)


def test_txt_blocks_unwrapped_and_blank_line_separated(txt_out):
    lines = txt_out.splitlines()
    # Each paragraph is one unwrapped line: the long paragraph 11 must be
    # a single line containing both its first and last words.
    hits = [l for l in lines
            if l.startswith("11. A party propounding interrogatories")]
    assert len(hits) == 1
    assert hits[0].endswith("informally before filing.")
    # Structure: blocks separated by blank lines.
    idx = lines.index(hits[0])
    assert lines[idx - 1] == "" and lines[idx + 1] == ""


def test_txt_footnotes_preserved_sensibly(txt_out):
    assert ": TKFN3" not in txt_out, "footnote def rendered as colon orphan"
    # Inline [n] markers survive on their referencing words, and the notes
    # land in an end-of-document notes section in number order.
    assert "here.[1]" in txt_out
    assert "period.[2]" in txt_out
    assert "size.[3]" in txt_out
    assert "\n[1] " in txt_out and "\n[2] " in txt_out and "\n[3] " in txt_out
    assert txt_out.find("TKFN3") > txt_out.find("TKFN2") > \
        txt_out.find("TKFN1") > txt_out.find("size.[3]")


def test_txt_declsignblock_keeps_declarant_name(txt_out):
    assert "JANE ROE" in txt_out
    assert "CaliforniaDeclarant" not in txt_out
    year = datetime.date.today().year
    assert (f"Executed this _____ day of _________________, {year}, "
            "at Springfield, California.") in txt_out
    assert "Declarant, In Pro Se" in txt_out
