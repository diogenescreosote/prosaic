#!/usr/bin/env python3
"""Convert a markdown pleading source to an editable .docx for proposed orders.

Reads the same YAML front matter + markdown body as md_pleading.py, but
produces a Word document suitable for the court to edit and sign. Includes
caption, double spacing, and proper margins, but omits line numbers and
the vertical rule (courts don't need those in editable documents).

Usage:
    python md_to_docx.py input.md output.docx
"""

from __future__ import annotations

import argparse
import datetime
import json
import re
import sys
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI

# The QR seam (payload resolution + qrencode rendering) lives in
# md_pleading so the three renderers cannot drift.
import md_pleading as mp
from docx.opc.constants import RELATIONSHIP_TYPE as RT, CONTENT_TYPE as CT
from xml.sax.saxutils import escape

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ---------------------------------------------------------------------------
# Layout constants — editable-Word rendition of California pleading style
# ---------------------------------------------------------------------------

BODY_FONT_NAME = "Times New Roman"
BODY_FONT_SIZE_PT = 12
BODY_LINE_SPACING_PT = 24        # exact double spacing for 12 pt body type
CAPTION_LINE_SPACING_PT = 14     # single-spaced caption / letter-header lines

# The `notreal:` banner, matching md_pleading's DRAFT_BANNER_* so the
# PDF and the DOCX of one source say the same thing the same way.
DRAFT_BANNER_FONT_SIZE_PT = 10
DRAFT_BANNER_RGB = (0xBF, 0x00, 0x00)
PAGE_MARGIN_TOP_IN = 1.0
PAGE_MARGIN_BOTTOM_IN = 1.0
PAGE_MARGIN_LEFT_IN = 1.0
PAGE_MARGIN_RIGHT_IN = 0.75
# Caption-table wrap heuristics (character count — Word wraps its own
# text, so these only decide how the caption table's rows are split; the
# PDF caption wraps by measured string width instead).
CAPTION_NAME_WRAP_CHARS = 30
CAPTION_TITLE_WRAP_CHARS = 35

# Shared front-end: everything up to "how output is emitted" lives in
# md_pleading so the DOCX and PDF renderers parse identically — front-matter
# parsing, variant resolution, redaction/date/posblock macros, exhibit
# citation resolution, inline formatting, flat auto-numbering, and footnote
# extraction. Only the docx emission below is renderer-specific.
from md_pleading import (
    SIG_KEEP_MAX_LEAD_BLOCKS,
    SIG_LEAD_MAX_CHARS,
    SUPPORTED_VARIANTS,
    REDACTION_SIDECAR_SUFFIX,
    Exhibit,
    alpha,
    apply_variant_to_meta,
    autonumber_list_items,
    draft_banner_text,
    effective_variant,
    extract_footnote_defs,
    load_external_exhibits,
    parse_front_matter,
    parse_inline_styles,
    roman,
    substitute_date_macro,
    substitute_exhibit_refs,
    substitute_posblock_macro,
    require_attachment_has_no_caption,
    warn_unknown_front_matter_keys,
    substitute_redaction_macros,
    suppresses_caption,
    unsigned_dated_line,
    unsigned_decl_execution_line,
)


def typographic_subs(text: str) -> str:
    # House style for DOCX: keep ASCII straight quotes and apostrophes; only
    # convert dashes to their proper glyphs.
    text = text.replace("---", "\u2014")
    text = text.replace("--", "\u2013")
    return text


def add_para(doc: Document, text: str, bold: bool = False,
             align=WD_ALIGN_PARAGRAPH.LEFT,
             spacing: float = BODY_LINE_SPACING_PT,
             font_size: float = BODY_FONT_SIZE_PT) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(spacing)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = align
    run = p.add_run(text)
    run.font.name = BODY_FONT_NAME
    run.font.size = Pt(font_size)
    run.bold = bold


def _add_draft_banner(section, meta: dict) -> None:
    """Put the `notreal:` banner in the page header, in red.

    The header is what makes it repeat on every page without touching
    the body, which matters because a DOCX goes to counsel to be edited
    — a banner inside the text would be something to delete by
    accident, or to leave in by accident. Clearing it means clearing
    `notreal:` in the source and rebuilding.
    """
    text = draft_banner_text(meta)
    if not text:
        return
    p = section.header.paragraphs[0]
    # Drop any existing runs outright. Assigning `.text = ""` would
    # leave an empty run in place, so the banner would land second and
    # inherit nothing.
    for existing in list(p.runs):
        existing._element.getparent().remove(existing._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(DRAFT_BANNER_FONT_SIZE_PT)
    run.font.color.rgb = RGBColor(*DRAFT_BANNER_RGB)


def add_blank(doc: Document) -> None:
    add_para(doc, "")


def add_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "6",
        qn("w:space"): "1", qn("w:color"): "000000",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)


def build_citation_exhibit_map(meta: dict, input_path: Path,
                               variant: str) -> dict:
    """Build a shortname -> Exhibit map for citation resolution only.

    The docx renderer never attaches exhibit files, so exhibit paths are
    neither resolved nor required to exist — only the letter (positional,
    like the PDF's) and any ``pages`` spec matter. ``exhibit_source``
    entries are appended after the file's own exhibits, in source order,
    so the lettering matches the PDF for the same document.
    """
    entries: List[Tuple[str, str | None]] = []
    for item in meta.get("exhibits") or []:
        if not isinstance(item, dict) or not item.get("shortname"):
            entries.append(("", None))  # placeholder keeps letters positional
            continue
        pages = item.get("pages")
        pages = str(pages).strip() if pages is not None else None
        entries.append((str(item["shortname"]).strip(), pages or None))

    exhibit_source = meta.get("exhibit_source")
    if exhibit_source:
        source_path = Path(exhibit_source)
        if not source_path.is_absolute():
            source_path = (input_path.parent / source_path).resolve()
        own = {sn for sn, _ in entries}
        for ex in load_external_exhibits(source_path, variant):
            if ex.shortname not in own:
                entries.append((ex.shortname, ex.pages))

    mapping: dict = {}
    for idx, (shortname, pages) in enumerate(entries, start=1):
        if shortname:
            mapping[shortname] = Exhibit(shortname=shortname, title="",
                                         path=None, letter=alpha(idx),
                                         pages=pages)
    return mapping


def _assign_footnote_numbers(body: str, footnote_defs: dict | None = None) -> dict:
    """Number footnote references in first-appearance (reading) order.

    Only ids with a definition get numbers: an undefined reference renders
    as its literal ``[^id]`` marker (see _add_spans) instead of emitting a
    ``w:footnoteReference`` to a note _attach_footnotes_part never creates
    (a dangling reference in the OPC package).
    """
    defs = footnote_defs or {}
    nums: dict = {}
    for m in re.finditer(r"\[\^([^\]]+?)\]", body):
        fid = m.group(1).strip()
        if fid in defs and fid not in nums:
            nums[fid] = len(nums) + 1
    return nums


def _format_letter_date_docx(raw: str) -> str:
    if raw.strip().lower() == "today":
        d = datetime.date.today()
        return f"{d.strftime('%B')} {d.day}, {d.year}"
    try:
        d = datetime.date.fromisoformat(raw.strip())
        return f"{d.strftime('%B')} {d.day}, {d.year}"
    except (ValueError, TypeError):
        return str(raw)


def _add_spans(p, text: str, base_bold: bool = False,
               force_italic: bool = False,
               font_size: float = BODY_FONT_SIZE_PT,
               footnote_numbers: dict | None = None) -> None:
    """Render inline markdown into runs on paragraph ``p`` using the shared
    ``parse_inline_styles`` front-end: bold, italic, underline, and superscript
    footnote-reference markers. Only the output (docx runs) is renderer-specific.
    """
    for span in parse_inline_styles(text):
        if span.footnote_id is not None:
            num = (footnote_numbers or {}).get(span.footnote_id)
            if num is None:
                run = p.add_run(f"[^{span.footnote_id}]")
                run.font.name = BODY_FONT_NAME
                run.font.size = Pt(font_size)
                continue
            # Emit a real footnote reference so Word/Pages render the note at
            # the page bottom and auto-number it (see _attach_footnotes_part).
            run = p.add_run()
            rpr = run._r.get_or_add_rPr()
            va = OxmlElement("w:vertAlign")
            va.set(qn("w:val"), "superscript")
            rpr.append(va)
            ref = OxmlElement("w:footnoteReference")
            ref.set(qn("w:id"), str(num))
            run._r.append(ref)
            continue
        if span.text == "":
            continue
        run = p.add_run(span.text)
        run.font.name = BODY_FONT_NAME
        run.font.size = Pt(font_size)
        run.bold = span.bold or base_bold
        run.italic = span.italic or force_italic
        if span.underline:
            run.underline = True
        if span.highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _add_styled_para(doc: Document, text: str, bold: bool = False,
                     align=WD_ALIGN_PARAGRAPH.LEFT,
                     spacing: float = BODY_LINE_SPACING_PT,
                     font_size: float = BODY_FONT_SIZE_PT,
                     force_italic: bool = False,
                     footnote_numbers: dict | None = None):
    """Add a paragraph, parsing shared inline emphasis (bold/italic/underline)
    and footnote markers. If bold=True, every run is forced bold (individual
    italic spans are preserved). Returns the paragraph for further manipulation.
    """
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(spacing)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = align
    _add_spans(p, text, base_bold=bold, force_italic=force_italic,
               font_size=font_size, footnote_numbers=footnote_numbers)
    return p


def _spans_to_run_xml(text: str) -> str:
    """Serialize inline-styled text into WordprocessingML <w:r> runs for use
    inside the footnotes part (which python-docx does not manage directly)."""
    runs: List[str] = []
    for span in parse_inline_styles(typographic_subs(text)):
        if span.footnote_id is not None:  # nested footnotes: render literally
            runs.append('<w:r><w:t xml:space="preserve">%s</w:t></w:r>'
                        % escape(f"[^{span.footnote_id}]"))
            continue
        if span.text == "":
            continue
        props = []
        if span.bold:
            props.append("<w:b/>")
        if span.italic:
            props.append("<w:i/>")
        if span.underline:
            props.append('<w:u w:val="single"/>')
        if span.highlight:
            props.append('<w:highlight w:val="yellow"/>')
        rpr = "<w:rPr>%s</w:rPr>" % "".join(props) if props else ""
        runs.append('<w:r>%s<w:t xml:space="preserve">%s</w:t></w:r>'
                    % (rpr, escape(span.text)))
    return "".join(runs)


def _attach_footnotes_part(doc: Document, footnote_numbers: dict | None,
                           footnote_defs: dict | None) -> None:
    """Add a real WordprocessingML footnotes part so footnotes render at the
    bottom of the page (Word/Pages auto-number them). Body references are
    emitted as <w:footnoteReference> runs by _add_spans.

    No third-party dependency: the footnotes part is injected into the OPC
    package directly via python-docx's part API (same output as native Word).
    """
    if not footnote_numbers or not footnote_defs:
        return
    items = sorted(
        (num, fid) for fid, num in footnote_numbers.items()
        if fid in footnote_defs
    )
    if not items:
        return
    entries = [
        '<w:footnote w:type="separator" w:id="-1"><w:p><w:r>'
        '<w:separator/></w:r></w:p></w:footnote>',
        '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r>'
        '<w:continuationSeparator/></w:r></w:p></w:footnote>',
    ]
    for num, fid in items:
        entries.append(
            '<w:footnote w:id="%d"><w:p>'
            '<w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
            '<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
            '<w:footnoteRef/></w:r>'
            '<w:r><w:t xml:space="preserve"> </w:t></w:r>%s'
            '</w:p></w:footnote>'
            % (num, _spans_to_run_xml(footnote_defs[fid]))
        )
    xml = ("<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
           '<w:footnotes xmlns:w="%s">%s</w:footnotes>'
           % (_W_NS, "".join(entries)))
    part = Part(PackURI("/word/footnotes.xml"), CT.WML_FOOTNOTES,
                xml.encode("utf-8"), doc.part.package)
    doc.part.relate_to(part, RT.FOOTNOTES)


# ---------------------------------------------------------------------------
# Word native multilevel numbering for legal headings (I. / A. / 1.)
# ---------------------------------------------------------------------------

_LEGAL_ABSTRACT_ID = "9100"
_LEGAL_NUM_ID = 9101


def _ensure_legal_numbering(doc: Document) -> int:
    """Ensure the document has a 3-level legal-outline multilevel numbering
    definition (I. / A. / 1.) and return the numId to apply to heading
    paragraphs via ``w:numPr``.

    Idempotent: safe to call multiple times; the numbering is only added
    once per document.
    """
    from docx.oxml import parse_xml
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    numbering_part = None
    try:
        numbering_part = doc.part.numbering_part
    except (KeyError, AttributeError):
        numbering_part = None

    if numbering_part is None:
        try:
            from docx.parts.numbering import NumberingPart
            numbering_part = NumberingPart.new()
            doc.part.relate_to(numbering_part, RT.NUMBERING)
        except Exception:
            # Fall back: build numbering XML from scratch and attach
            import os
            from docx.opc.part import PartFactory
            from docx.opc.packuri import PackURI
            default_xml = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'''
            numbering_part = PartFactory.part_class_selector(
                RT.NUMBERING, "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"
            )(
                PackURI("/word/numbering.xml"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
                default_xml, doc.part.package,
            )
            doc.part.relate_to(numbering_part, RT.NUMBERING)

    numbering_elem = numbering_part.element

    # Check if already defined
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for abs_num in numbering_elem.findall(qn("w:abstractNum")):
        if abs_num.get(qn("w:abstractNumId")) == _LEGAL_ABSTRACT_ID:
            return _LEGAL_NUM_ID

    abstract_xml = f'''<w:abstractNum xmlns:w="{w}" w:abstractNumId="{_LEGAL_ABSTRACT_ID}">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0">
    <w:start w:val="1"/>
    <w:numFmt w:val="upperRoman"/>
    <w:suff w:val="space"/>
    <w:lvlText w:val="%1."/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="0" w:firstLine="0"/>
    </w:pPr>
  </w:lvl>
  <w:lvl w:ilvl="1">
    <w:start w:val="1"/>
    <w:numFmt w:val="upperLetter"/>
    <w:suff w:val="space"/>
    <w:lvlText w:val="%2."/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="360" w:firstLine="0"/>
    </w:pPr>
  </w:lvl>
  <w:lvl w:ilvl="2">
    <w:start w:val="1"/>
    <w:numFmt w:val="decimal"/>
    <w:suff w:val="space"/>
    <w:lvlText w:val="%3."/>
    <w:lvlJc w:val="left"/>
    <w:pPr>
      <w:ind w:left="720" w:firstLine="0"/>
    </w:pPr>
  </w:lvl>
</w:abstractNum>'''
    abstract_elem = parse_xml(abstract_xml)

    first_num = numbering_elem.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract_elem)
    else:
        numbering_elem.append(abstract_elem)

    num_xml = f'''<w:num xmlns:w="{w}" w:numId="{_LEGAL_NUM_ID}">
  <w:abstractNumId w:val="{_LEGAL_ABSTRACT_ID}"/>
</w:num>'''
    num_elem = parse_xml(num_xml)
    numbering_elem.append(num_elem)

    return _LEGAL_NUM_ID


def _apply_heading_numbering(paragraph, num_id: int, ilvl: int) -> None:
    """Attach w:numPr to the paragraph so Word auto-numbers it under the
    given multilevel list numId at the specified ilvl (0, 1, or 2)."""
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()

    # Remove any existing numPr so we don't double up on re-application.
    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), str(ilvl))
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_elem)
    numPr.append(numId_elem)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# Signature block rendering
# ---------------------------------------------------------------------------

def _keep_block_together(doc: Document, start_index: int) -> None:
    """Keep every paragraph of a just-emitted block (paragraphs from
    ``start_index`` onward) with the next, so Word never splits the block
    across a page break."""
    paras = doc.paragraphs[start_index:]
    for p in paras[:-1]:
        p.paragraph_format.keep_with_next = True


def _attach_lead(doc: Document,
                 max_leads: int = SIG_KEEP_MAX_LEAD_BLOCKS) -> None:
    """Keep the short lead-in paragraph(s) immediately preceding a signature
    block on the same page as the block. This is the default behavior — a
    heading or short lead (e.g., "ORDER", "IT IS SO ORDERED.") is never
    stranded above the signature it introduces. Long paragraphs are not
    treated as leads, so ordinary body text is unaffected.

    The policy (how many leads, how short is "short") is the shared
    keep-with-signature policy defined in md_pleading; the PDF emitter
    measures grid lines, this emitter uses the character-count proxy."""
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    for p in reversed(non_empty[-max_leads:]):
        if len(p.text.strip()) <= SIG_LEAD_MAX_CHARS:
            p.paragraph_format.keep_with_next = True
        else:
            break


def _emit_notarial(doc: Document, block) -> None:
    """A California notarial certificate as a bordered one-cell table,
    sans-serif, mirroring the PDF object (statutory wording from
    md_pleading's constants so the renderers cannot drift)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""

    def para(text, bold=False, center=False, size=9.5):
        pr = cell.add_paragraph()
        if center:
            pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pr.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        return pr

    para(mp.NOTARIAL_TITLES[block.kind], bold=True, center=True, size=10.5)
    para(mp.NOTARIAL_DISCLOSURE, size=8.5)
    para("State of California")
    para("County of _____________________________ )")
    para(mp.notarial_text(block))
    if block.kind == "acknowledgment":
        para(mp.ACK_PERJURY)
    if block.kind in ("acknowledgment", "proofexec"):
        para(mp.ACK_WITNESS_LINE)
    for _ in range(5):  # clear zone for the seal
        para("")
    para("Signature _________________________________          (Seal)")
    add_blank(doc)


def _emit_witnessattest(doc: Document, names_raw: str,
                        esign_role_counter: list | None = None) -> None:
    for name in [n.strip() for n in names_raw.split("\\\\") if n.strip()]:
        add_blank(doc)
        p1 = doc.add_paragraph()
        p1.add_run("____________________________      Date: ______________")
        if esign_role_counter is not None:
            esign_role_counter[0] += 1
            n = esign_role_counter[0]
            _esign_tag_para(doc, "{{Signature %d;role=Signer %d;type=signature}} "
                                 "{{Date %d;role=Signer %d;type=date}}"
                            % ((n,) * 4))
        cap = doc.add_paragraph()
        cap.add_run(f"Signature of {name}").font.size = Pt(9)
        p2 = doc.add_paragraph()
        p2.add_run("____________________________")
        if esign_role_counter is not None:
            _esign_tag_para(doc, "{{Residence %d;role=Signer %d;type=text}}"
                            % (esign_role_counter[0], esign_role_counter[0]))
        cap2 = doc.add_paragraph()
        cap2.add_run("Residing at (city and state)").font.size = Pt(9)


def _emit_barcode(doc: Document, fmt: str, payload: str, caption: str) -> None:
    """Embed a barcode image at its printed size (md_pleading computes
    the module geometry), capped at 6.5 inches of column width."""
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".png") as tmp:
        w_pt, h_pt = mp.barcode_image(fmt, payload, tmp.name)
        w_in = min(w_pt / 72.0, 6.5)
        doc.add_picture(tmp.name, width=Inches(w_in))
    if caption:
        para = doc.add_paragraph()
        para.add_run(caption).font.size = Pt(12)
    add_blank(doc)


def _emit_leftright(doc: Document, left: str, right: str) -> None:
    """Left text at the margin, right text flush right via a
    right-aligned tab stop at the text edge."""
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches as _In

    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.tab_stops.add_tab_stop(
        _In(6.5), WD_TAB_ALIGNMENT.RIGHT)
    _add_spans(para, left)
    para.add_run("\t")
    _add_spans(para, right)


def _emit_sigrow_docx(doc: Document, left_raw: str, right_raw: str) -> None:
    left = [s.strip() for s in left_raw.split("\\\\")]
    right = [s.strip() for s in right_raw.split("\\\\")]
    add_blank(doc)
    para = doc.add_paragraph()
    para.add_run("_" * 34 + "        " + "_" * 34)
    for i in range(max(len(left), len(right))):
        row = doc.add_paragraph()
        lft = left[i] if i < len(left) else ""
        rgt = right[i] if i < len(right) else ""
        row.add_run(f"{lft:<44}{rgt}")
    add_blank(doc)


def _emit_fixedwidth(doc: Document, text: str) -> None:
    """Verbatim monospace lines; no typographic substitutions."""
    for raw in text.split("\n"):
        para = doc.add_paragraph()
        run = para.add_run(raw)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    add_blank(doc)


def _esign_tag_para(doc: Document, tag: str) -> None:
    """A DocuSeal text tag in white 6 pt: invisible in print, present
    in the document text for field placement, stripped from the
    executed document by remove_tags. The PDF is the primary e-sign
    artifact; DOCX tags land adjacent to (not atop) their blanks."""
    para = doc.add_paragraph()
    run = para.add_run(tag)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _emit_whereofsignblock(doc: Document, name: str, role: str,
                           instrument: str, esign_role: int) -> None:
    start = len(doc.paragraphs)
    add_blank(doc)
    _add_styled_para(doc, mp.WHEREOF_CLAUSE.format(
        name=name, instrument=instrument or "instrument"))
    _esign_tag_para(doc, "{{Day %d;role=Signer %d;type=text}} "
                         "{{Month %d;role=Signer %d;type=text}} "
                         "{{Year %d;role=Signer %d;type=text}} "
                         "{{Location %d;role=Signer %d;type=text}}"
                    % ((esign_role,) * 8))
    add_blank(doc)
    _add_styled_para(doc, "____________________________________")
    _esign_tag_para(doc, "{{Signature %d;role=Signer %d;type=signature}}"
                    % (esign_role, esign_role))
    add_blank(doc)
    _add_styled_para(doc, name)
    if role:
        _add_styled_para(doc, role)
    _keep_block_together(doc, start)


def _emit_signblock(doc: Document, name: str, role: str) -> None:
    """Emit a standard signature block: Dated line / blank / signature line
    / blank / NAME / Role."""
    year = datetime.date.today().year
    start = len(doc.paragraphs)
    _add_styled_para(doc, unsigned_dated_line(year))
    add_blank(doc)
    _add_styled_para(doc, "____________________________________")
    add_blank(doc)
    _add_styled_para(doc, name)
    if role:
        _add_styled_para(doc, role)
    _keep_block_together(doc, start)
    add_blank(doc)


def _emit_declsignblock(doc: Document, name: str, location: str,
                        role: str) -> None:
    """Emit a declaration signature block: Executed this ___ day of ____,
    YEAR, at LOCATION / blank / signature line / blank / NAME / Role."""
    year = datetime.date.today().year
    start = len(doc.paragraphs)
    _add_styled_para(doc, unsigned_decl_execution_line(year, location))
    add_blank(doc)
    _add_styled_para(doc, "____________________________________")
    add_blank(doc)
    _add_styled_para(doc, name)
    if role:
        _add_styled_para(doc, role)
    _keep_block_together(doc, start)
    add_blank(doc)


# ---------------------------------------------------------------------------
# Signblock argument parser
# ---------------------------------------------------------------------------

_SIGNBLOCK_RE = re.compile(r"^\\signblock((?:\{[^{}]*\})+)\s*$")
_DECLSIGNBLOCK_RE = re.compile(
    r"^\\declsignblock\{(.+?)\}\{(.+?)\}(?:\{(.*?)\})?\s*$"
)
_JUDGESIGNBLOCK_RE = re.compile(r"^\\judgesignblock\{(.+?)\}\s*$")
_ACK_RE = re.compile(r"^\\acknowledgment\{(.*?)\}\s*$")
_JURAT_RE = re.compile(r"^\\jurat\{(.*?)\}\s*$")
_PROOF_RE = re.compile(r"^\\proofofexecution\{(.*?)\}\{(.*?)\}\s*$")
_WATT_RE = re.compile(r"^\\witnessattestation\{(.+?)\}\s*$")
_LEFTRIGHT_RE = re.compile(r"^\\leftright\{(.*?)\}\{(.*?)\}\s*$")
_CENTER_RE = re.compile(r"^\\center\{(.*?)\}\s*$")
_SIGROW_RE = re.compile(r"^\\sigrow\{(.*?)\}\{(.*?)\}\s*$")
_BARCODE_RE = re.compile(r"^\\barcode\{([a-z0-9]+)\}\{(.+?)\}(?:\{(.*?)\})?\s*$")
_BARCODEFILE_RE = re.compile(r"^\\barcodefile\{([a-z0-9]+)\}\{(.+?)\}(?:\{(.*?)\})?\s*$")


def build_letter_header(doc: Document, meta: dict) -> None:
    # Header metadata is YAML and bypasses the markdown pipeline; route it
    # through the docx typographic_subs so dashes convert exactly as they
    # do in the body (mirrors the PDF's _letter_header).
    CL = CAPTION_LINE_SPACING_PT
    from_lines = [
        meta["filer_name"],
        *[str(a) for a in meta["filer_address_lines"]],
    ]
    if meta.get("filer_phone"):
        from_lines.append(meta["filer_phone"])
    if meta.get("filer_email"):
        from_lines.append(meta["filer_email"])
    for line in from_lines:
        add_para(doc, typographic_subs(str(line)),
                 align=WD_ALIGN_PARAGRAPH.RIGHT, spacing=CL)

    add_blank(doc)

    to_lines = [
        meta["to_name"],
        *[str(a) for a in meta["to_address_lines"]],
    ]
    for line in to_lines:
        add_para(doc, typographic_subs(str(line)), spacing=CL)

    add_blank(doc)

    date_str = _format_letter_date_docx(meta.get("date", "_______________"))
    add_para(doc, date_str, spacing=CL)

    service_method = meta.get("service_method")
    if service_method:
        add_para(doc, str(service_method), bold=True, spacing=CL)

    add_blank(doc)

    add_para(doc, "Re: " + typographic_subs(str(meta["paper_title"])),
             bold=True, spacing=CL)
    add_blank(doc)


def build_body_letter(doc: Document, body: str, meta: dict,
                      footnote_numbers: dict | None = None,
                      footnote_defs: dict | None = None) -> None:
    body = _COMMENT_RE.sub("", body)
    body = typographic_subs(body)

    h1 = h2 = h3 = 0
    for para_text in body.split("\n\n"):
        text = " ".join(para_text.split())
        if not text:
            continue

        m_sign = re.match(r"^\\lettersignblock\{(.+)\}$", text)
        if m_sign:
            add_blank(doc)
            _add_styled_para(doc, "Sincerely,")
            add_blank(doc)
            add_blank(doc)
            for line in m_sign.group(1).split("\\n"):
                _add_styled_para(doc, line.strip())
            continue

        m_heading = re.match(r"^(#{1,3})\s+(.*)$", text)
        if m_heading:
            level = len(m_heading.group(1))
            heading_text = m_heading.group(2)
            if level == 1:
                h1 += 1; h2 = h3 = 0
                prefix = f"{roman(h1)}."
            elif level == 2:
                h2 += 1; h3 = 0
                prefix = f"{alpha(h2)}."
            else:
                h3 += 1
                prefix = f"{h3}."
            # heading_numbers: false — headings that carry their own
            # enumeration ("Article I. ...") skip the auto prefix.
            if not meta.get("heading_numbers", True):
                prefix = ""
            add_blank(doc)
            _add_styled_para(doc, f"{prefix} {heading_text}".strip(), bold=True,
                             footnote_numbers=footnote_numbers)
            continue

        _add_styled_para(doc, text, footnote_numbers=footnote_numbers)


def build_caption(doc: Document, meta: dict) -> None:
    # Caption metadata is YAML and bypasses the markdown pipeline; route it
    # through the docx typographic_subs so dashes convert exactly as they do
    # in the body (mirrors the PDF caption; spec: pleading_markdown_spec.md,
    # typographic substitutions apply to YAML display strings).
    CL = CAPTION_LINE_SPACING_PT

    for line in [meta["filer_name"],
                 *meta["filer_address_lines"],
                 meta["filer_phone"],
                 meta["filer_email"]]:
        add_para(doc, typographic_subs(str(line)), spacing=CL)
    add_para(doc, typographic_subs(meta["filer_role"]).upper(), spacing=CL)
    add_blank(doc)

    add_para(doc, typographic_subs(meta["court_name"]), bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing=CL)
    add_para(doc, typographic_subs(meta["court_county"]), bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, spacing=CL)
    add_blank(doc)

    pet = typographic_subs(meta["petitioner"])
    resp = typographic_subs(meta["respondent"])
    cap_first = meta.get("caption_first_party_label", "Petitioner")
    cap_second = meta.get("caption_second_party_label", "Respondent")
    case_num = typographic_subs(str(meta.get("case_number", "_______________")))
    title = typographic_subs(meta["paper_title"]).upper()

    table = doc.add_table(rows=0, cols=3)
    table.columns[0].width = Inches(2.75)
    table.columns[1].width = Inches(0.25)
    table.columns[2].width = Inches(3.5)

    rows_data: List[Tuple[str, str, str, bool, bool]] = []

    for wrapped in _wrap_name(pet):
        rows_data.append((wrapped, ")", case_num if not rows_data else "", True, False))
    rows_data.append(("", ")", "", False, False))
    rows_data.append((f"     {cap_first},", ")", "", False, False))

    title_lines = _wrap_title(title)
    for i, tl in enumerate(title_lines):
        left = "" if i > 0 else ""
        rows_data.append((left, ")", tl, False, True))

    rows_data.append(("     vs.", ")", "", False, False))
    rows_data.append(("", ")", "", False, False))

    for wrapped in _wrap_name(resp):
        rows_data.append((wrapped, ")", "", True, False))
    rows_data.append(("", ")", "", False, False))
    rows_data.append((f"     {cap_second}.", ")", "", False, False))

    for left, mid, right, left_bold, right_bold in rows_data:
        row = table.add_row()
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(CL)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)

        run = row.cells[0].paragraphs[0].add_run(left)
        run.font.size = Pt(BODY_FONT_SIZE_PT)
        run.font.name = BODY_FONT_NAME
        run.bold = left_bold

        run = row.cells[1].paragraphs[0].add_run(mid)
        run.font.size = Pt(BODY_FONT_SIZE_PT)
        run.font.name = BODY_FONT_NAME

        run = row.cells[2].paragraphs[0].add_run(right)
        run.font.size = Pt(BODY_FONT_SIZE_PT)
        run.font.name = BODY_FONT_NAME
        run.bold = right_bold

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none", qn("w:sz"): "0",
            qn("w:space"): "0", qn("w:color"): "auto",
        })
        borders.append(el)
    tblPr.append(borders)

    add_rule(doc)
    add_blank(doc)


def _wrap_name(name: str, max_len: int = CAPTION_NAME_WRAP_CHARS) -> List[str]:
    if len(name) <= max_len:
        return [name]
    words = name.split()
    lines, current = [], ""
    for w in words:
        trial = f"{current} {w}".strip()
        if len(trial) <= max_len:
            current = trial
        else:
            if current:
                lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines


def _wrap_title(title: str, max_len: int = CAPTION_TITLE_WRAP_CHARS) -> List[str]:
    if len(title) <= max_len:
        return [title]
    words = title.split()
    lines, current = [], ""
    for w in words:
        trial = f"{current} {w}".strip()
        if len(trial) <= max_len:
            current = trial
        else:
            if current:
                lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines


def build_body(doc: Document, body: str, meta: dict,
               footnote_numbers: dict | None = None,
               footnote_defs: dict | None = None) -> None:
    """Render the markdown body into the document with:
      * inline **bold**, *italic*, ***bold italic***, <u>underline</u> parsing
      * [^id] footnote references (superscript) + notes emitted at the end
      * \\signblock / \\declsignblock / \\judgesignblock macro expansion
      * # / ## / ### headings with Word-native multilevel auto-numbering
    (\\exhibit{}/\\attachment{} references are already resolved by main()
    via the shared substitute_exhibit_refs before this runs)
    """
    # Strip HTML comments first (they're strategic annotations not meant for
    # the rendered output). Exhibit references were already resolved in
    # main() via the shared substitute_exhibit_refs.
    body = _COMMENT_RE.sub("", body)

    # \fixedwidth{ ... } content is verbatim: stash it BEFORE the
    # body-wide typographic pass, which would otherwise turn an armored
    # block's ----- runs into em dashes (corruption, not style).
    fw_blocks: list[str] = []

    def _fw_stash(m: re.Match) -> str:
        fw_blocks.append(m.group(1))
        return f"\n\n@@FIXEDWIDTH{len(fw_blocks) - 1}@@\n\n"

    body = re.sub(r"(?ms)^\\fixedwidth\{[ \t]*\n(.*?)\n\}[ \t]*$",
                  _fw_stash, body)

    body = typographic_subs(body)

    num_id = None  # lazy-initialized the first time a heading is seen
    esign_role_counter = [0]  # DocuSeal signer roles, document order

    paragraphs = body.split("\n\n")
    for para_text in paragraphs:
        stripped = para_text.strip()
        if not stripped:
            continue

        # Preserve inter-paragraph whitespace collapse but keep the line
        # content joined. This matches md_pleading.py's body layout.
        text = " ".join(stripped.split())

        # Signature blocks ------------------------------------------------
        m = _DECLSIGNBLOCK_RE.match(text)
        if m:
            name = m.group(1).strip()
            location = m.group(2).strip()
            role_override = (m.group(3) or "").strip()
            role = role_override if role_override else str(
                meta.get("filer_role", "")
            ).title()
            _attach_lead(doc)
            _emit_declsignblock(doc, name, location, role)
            esign_role_counter[0] += 1
            _esign_tag_para(doc, "{{Date %d;role=Signer %d;type=text}} "
                                 "{{Signature %d;role=Signer %d;type=signature}}"
                            % ((esign_role_counter[0],) * 4))
            continue

        m = _ACK_RE.match(text)
        if m:
            _attach_lead(doc)
            _emit_notarial(doc, mp.Block("acknowledgment", m.group(1).strip()))
            continue
        m = _JURAT_RE.match(text)
        if m:
            _attach_lead(doc)
            _emit_notarial(doc, mp.Block("jurat", m.group(1).strip()))
            continue
        m = _PROOF_RE.match(text)
        if m:
            _attach_lead(doc)
            _emit_notarial(doc, mp.Block("proofexec", m.group(1).strip(),
                                         spans=[mp.TextSpan(m.group(2).strip())]))
            continue
        m = _WATT_RE.match(text)
        if m:
            _attach_lead(doc)
            _emit_witnessattest(doc, m.group(1).strip(), esign_role_counter)
            continue
        m = _LEFTRIGHT_RE.match(text)
        if m:
            _emit_leftright(doc, m.group(1).strip(), m.group(2).strip())
            continue
        m = _CENTER_RE.match(text)
        if m:
            para = _add_styled_para(doc, m.group(1).strip())
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        m = _SIGROW_RE.match(text)
        if m:
            _attach_lead(doc)
            _emit_sigrow_docx(doc, m.group(1).strip(), m.group(2).strip())
            continue
        m = re.match(r"^@@FIXEDWIDTH(\d+)@@$", text)
        if m:
            _emit_fixedwidth(doc, fw_blocks[int(m.group(1))])
            continue

        m = _BARCODE_RE.match(text)
        if m:
            _attach_lead(doc)
            _emit_barcode(doc, m.group(1).strip(), m.group(2).strip(),
                          (m.group(3) or "").strip())
            continue

        m = _BARCODEFILE_RE.match(text)
        if m:
            block = mp.Block("barcodefile", m.group(2).strip(),
                             spans=[mp.TextSpan(m.group(1).strip())])
            _attach_lead(doc)
            _emit_barcode(doc, m.group(1).strip(), mp.barcode_payload(block),
                          (m.group(3) or "").strip())
            continue

        m = _SIGNBLOCK_RE.match(text)
        if m:
            args = [a.strip() for a in re.findall(r"\{([^{}]*)\}", m.group(1))]
            style = args[0] if args and args[0] in mp.SIGNBLOCK_STYLES else None
            _attach_lead(doc)
            if style == "dated":
                role = (args[2].title() if len(args) > 2 and args[2]
                        else str(meta.get("filer_role", "")).title())
                _emit_signblock(doc, args[1] if len(args) > 1 else "", role)
                esign_role_counter[0] += 1
                _esign_tag_para(doc, "{{Date %d;role=Signer %d;type=text}} "
                                     "{{Signature %d;role=Signer %d;type=signature}}"
                                % ((esign_role_counter[0],) * 4))
            elif style == "decl":
                role = (args[3] if len(args) > 3 and args[3]
                        else str(meta.get("filer_role", "")).title())
                _emit_declsignblock(doc, args[1] if len(args) > 1 else "",
                                    args[2] if len(args) > 2 else "", role)
                esign_role_counter[0] += 1
                _esign_tag_para(doc, "{{Date %d;role=Signer %d;type=text}} "
                                     "{{Signature %d;role=Signer %d;type=signature}}"
                                % ((esign_role_counter[0],) * 4))
            elif style == "judge":
                title = args[1] if len(args) > 1 else ""
                start = len(doc.paragraphs)
                add_blank(doc)
                _add_styled_para(doc, "Dated: _________________")
                add_blank(doc)
                _add_styled_para(doc, "____________________________________")
                add_blank(doc)
                _add_styled_para(doc, title)
                _keep_block_together(doc, start)
            elif style == "letter":
                add_blank(doc)
                _add_styled_para(doc, "Sincerely,")
                add_blank(doc)
                add_blank(doc)
                for line in (args[1] if len(args) > 1 else "").split("\\\\"):
                    _add_styled_para(doc, line.strip())
            elif style == "whereof":
                esign_role_counter[0] += 1
                _emit_whereofsignblock(
                    doc, args[1] if len(args) > 1 else "",
                    args[2] if len(args) > 2 else "",
                    args[3] if len(args) > 3 else "",
                    esign_role_counter[0])
            else:
                print("WARNING: \\signblock legacy form; write "
                      "\\signblock{dated}{NAME}{ROLE}", file=sys.stderr)
                role = (args[1].title() if len(args) > 1 and args[1]
                        else str(meta.get("filer_role", "")).title())
                _emit_signblock(doc, args[0] if args else "", role)
            continue

        m = _JUDGESIGNBLOCK_RE.match(text)
        if m:
            # Judge signature: Dated line, signature, role. Keep it together and
            # attached to the lead-in (e.g., "ORDER" / "IT IS SO ORDERED.").
            title = m.group(1).strip()
            _attach_lead(doc)
            start = len(doc.paragraphs)
            add_blank(doc)
            _add_styled_para(doc, "Dated: _________________")
            add_blank(doc)
            _add_styled_para(doc, "____________________________________")
            add_blank(doc)
            _add_styled_para(doc, title)
            _keep_block_together(doc, start)
            continue

        # Headings --------------------------------------------------------
        m = re.match(r"^(#{1,3})\s+(.*)$", text)
        if m:
            level = len(m.group(1))  # 1, 2, or 3
            heading_text = m.group(2).strip()
            if num_id is None:
                num_id = _ensure_legal_numbering(doc)
            add_blank(doc)
            # Emit the heading text (bold). Word fills in the I. / A. / 1.
            # via the multilevel numbering definition applied below.
            p = _add_styled_para(doc, heading_text, bold=True,
                                 footnote_numbers=footnote_numbers)
            _apply_heading_numbering(p, num_id, level - 1)
            p.paragraph_format.keep_with_next = True  # never orphan a heading
            continue

        # Block quotes --------------------------------------------------------
        if stripped.startswith(">"):
            # Join block-quote lines; preserve the quote-as-block.
            cleaned = " ".join(
                line.lstrip(">").strip() for line in stripped.splitlines()
            )
            p = _add_styled_para(doc, cleaned, footnote_numbers=footnote_numbers)
            p.paragraph_format.left_indent = Inches(0.5)
            continue

        # Bullet lists --------------------------------------------------------
        if stripped.startswith("- ") or stripped.startswith("* "):
            # Treat consecutive bullet lines within this paragraph chunk as
            # individual list items. The md_pleading.py convention only
            # supports one-level bullets.
            for line in stripped.splitlines():
                line = line.strip()
                if line.startswith("- ") or line.startswith("* "):
                    item = line[2:].strip()
                    p = _add_styled_para(doc, item,
                                         footnote_numbers=footnote_numbers)
                    p.style = doc.styles["List Bullet"]
            continue

        # Plain paragraph -------------------------------------------------
        _add_styled_para(doc, text, footnote_numbers=footnote_numbers)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert markdown pleading to editable .docx")
    parser.add_argument("input", help="Markdown source file")
    parser.add_argument("output", help="Output .docx path")
    parser.add_argument("--variant", choices=sorted(SUPPORTED_VARIANTS),
                        default=None,
                        help="Render variant to use for variant-aware metadata "
                             "and redactions. When omitted, redaction-bearing "
                             "documents build the PUBLIC variant (with a "
                             "warning); see md_pleading.effective_variant.")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found", file=sys.stderr)
        sys.exit(1)

    raw = input_path.read_text(encoding="utf-8")
    meta, body = parse_front_matter(raw)
    variant = effective_variant(meta, body, args.variant)
    if args.variant is None and variant == "public":
        print(
            f"WARNING: {input_path.name} carries redactions or "
            "variant-aware content but no --variant was given; building "
            "the PUBLIC (redacted) variant. Pass --variant sealed "
            "explicitly to render sealed content.",
            file=sys.stderr,
            flush=True,
        )
    meta = apply_variant_to_meta(meta, variant)

    # Exhibit citations resolve through the shared substitute_exhibit_refs
    # (both \exhibit{} and \attachment{}, doctype-aware label, en-dash page
    # cites) against a letters-only map — the docx attaches no files.
    doctype = meta.get("doctype", "pleading")
    exhibit_map = build_citation_exhibit_map(meta, input_path, variant)
    body = substitute_redaction_macros(body, meta, variant)
    body = substitute_posblock_macro(body, meta)
    body = substitute_exhibit_refs(body, exhibit_map, doctype=doctype)
    body = substitute_date_macro(body, meta)

    # Shared front-end transforms (identical to the PDF renderer): expand the
    # '#.' flat auto-numbering sentinel, pull out footnote definitions, and
    # number footnote references in reading order.
    body = autonumber_list_items(body)
    body, footnote_defs = extract_footnote_defs(body)
    footnote_numbers = _assign_footnote_numbers(body, footnote_defs)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(PAGE_MARGIN_TOP_IN)
        section.bottom_margin = Inches(PAGE_MARGIN_BOTTOM_IN)
        section.left_margin = Inches(PAGE_MARGIN_LEFT_IN)
        section.right_margin = Inches(PAGE_MARGIN_RIGHT_IN)
        _add_draft_banner(section, meta)

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT_NAME
    style.font.size = Pt(BODY_FONT_SIZE_PT)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)

    warn_unknown_front_matter_keys(meta, Path(args.input).name)
    require_attachment_has_no_caption(meta, Path(args.input).name)

    is_letter = doctype == "letter"
    if is_letter:
        build_letter_header(doc, meta)
        build_body_letter(doc, body, meta, footnote_numbers, footnote_defs)
    elif doctype == "document":
        # A plain document (contract, estate instrument): centered bold
        # title, no caption -- mirrors the PDF renderer's document mode.
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(str(meta.get("paper_title", "")).upper())
        run.bold = True
        add_blank(doc)
        build_body(doc, body, meta, footnote_numbers, footnote_defs)
    else:
        # An attachment to a JC form continues that form and must not
        # reintroduce the caption. The PDF renderer has honored this
        # since `no_caption:` existed; this one did not, so the same
        # source produced a correct PDF and a captioned DOCX.
        if not suppresses_caption(meta):
            build_caption(doc, meta)
        build_body(doc, body, meta, footnote_numbers, footnote_defs)

    _attach_footnotes_part(doc, footnote_numbers, footnote_defs)

    out_path = Path(args.output)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    # The sidecar quotes the verbatim sealed text, so it accompanies only
    # the sealed variant — a public/unscoped output directory may ship as
    # "the public packet" and must stay free of sealed bytes (mirrors
    # md_pleading.main).
    redaction_log = meta.get("_redaction_log") or []
    log_path = out_path.with_suffix(out_path.suffix + REDACTION_SIDECAR_SUFFIX)
    if redaction_log and variant == "sealed":
        with open(log_path, "w", encoding="utf-8") as f:
            json.dump(redaction_log, f, indent=2, ensure_ascii=True)
            f.write("\n")
        print(f"Wrote {log_path}")
    elif log_path.exists():
        # Never leave a stale sidecar from an earlier build next to a
        # current output that no longer warrants one.
        log_path.unlink()
        print(f"Removed stale {log_path}")


if __name__ == "__main__":
    main()
